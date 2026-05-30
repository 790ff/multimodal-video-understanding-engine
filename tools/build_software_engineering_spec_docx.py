from __future__ import annotations

from collections import defaultdict, deque
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "Software_Engineering_Specification.md"
OUTPUT = ROOT / "docs" / "Multimodal_Video_Understanding_Engine_SES_v0.6.docx"

DIAGRAM_IMAGES = {
    "B.2 Use Case Diagram": "use_case_diagram.png",
    "B.4 System Context Diagram": "system_context_diagram.png",
    "B.5 Level-1 Data Flow Diagram": "level_1_dfd.png",
    "B.6 Processing Pipeline Diagram": "processing_pipeline_diagram.png",
    "B.7 Activity Diagram": "activity_diagram.png",
    "B.8 Upload and Analysis Sequence Diagram": "upload_analysis_sequence.png",
    "B.9 Ask Video Sequence Diagram": "ask_video_sequence.png",
    "B.10 Domain Class Diagram": "domain_class_diagram.png",
    "B.11 Entity Relationship Diagram": "erd.png",
    "B.12 Video Status State Diagram": "video_status_state_diagram.png",
    "B.13 Component Diagram": "component_diagram_clean.png",
    "B.14 Deployment Diagram": "deployment_diagram.png",
    "B.15 API and Integration Swimlane Diagram": "api_integration_swimlane.png",
}


ACCENT = "111827"
MUTED = "4B5563"
LIGHT_FILL = "F3F4F6"
SOFT_FILL = "F9FAFB"


def slugify_bookmark(text: str, index: int) -> str:
    slug = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    if not slug:
        slug = "section"
    return f"toc_{index}_{slug[:28]}"


def collect_toc_entries(lines: list[str]) -> list[dict[str, str | int]]:
    entries: list[dict[str, str | int]] = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("## ") and stripped != "## Table of Contents":
            text = stripped[3:].strip()
            entries.append({"level": 1, "text": text})
        elif stripped.startswith("### "):
            text = stripped[4:].strip()
            entries.append({"level": 2, "text": text})

    for index, entry in enumerate(entries, start=1):
        entry["anchor"] = slugify_bookmark(str(entry["text"]), index)
        entry["bookmark_id"] = index
    return entries


def collect_table_entries(lines: list[str]) -> dict[int, dict[str, str | int]]:
    entries: dict[int, dict[str, str | int]] = {}
    current_heading = "Document Table"
    number = 1
    skip_headings = {"Table of Contents", "List of Tables", "List of Figures"}
    for index, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            if heading not in skip_headings:
                current_heading = heading
        elif stripped.startswith("### "):
            current_heading = stripped[4:].strip()
        elif stripped.startswith("#### "):
            current_heading = stripped[5:].strip()

        if stripped.startswith("|") and (index == 0 or not lines[index - 1].strip().startswith("|")):
            entries[index] = {
                "number": number,
                "title": current_heading,
                "anchor": slugify_bookmark(f"table_{number}_{current_heading}", number),
                "bookmark_id": 2000 + number,
            }
            number += 1
    return entries


def collect_figure_entries(lines: list[str]) -> dict[int, dict[str, str | int]]:
    entries: dict[int, dict[str, str | int]] = {}
    number = 1
    for index, line in enumerate(lines):
        stripped = line.strip()
        match = re.fullmatch(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)", stripped)
        if match:
            title = match.group("alt").strip() or "Figure"
            entries[index] = {
                "number": number,
                "title": title,
                "anchor": slugify_bookmark(f"figure_{number}_{title}", 3000 + number),
                "bookmark_id": 3000 + number,
            }
            number += 1
    return entries


def build_heading_queue(entries: list[dict[str, str | int]]):
    queue = defaultdict(deque)
    for entry in entries:
        queue[(entry["level"], entry["text"])].append(entry)
    return queue


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    r_pr.append(color)

    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_external_hyperlink(paragraph, text: str, url: str, underline: bool = True) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    r_pr.append(color)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, instruction: str, placeholder: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instr_run._r.append(instr)

    sep_run = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(sep)

    text_run = paragraph.add_run(placeholder)
    text_run.font.color.rgb = RGBColor.from_string(ACCENT)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_reference_list(doc: Document, entries: list[dict[str, str | int]], label: str) -> None:
    for entry in entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.left_indent = Inches(0.22 if entry.get("level") == 2 else 0)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.45), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        add_internal_hyperlink(paragraph, str(entry["text"]), str(entry["anchor"]))
        paragraph.add_run("\t")
        add_field(paragraph, f'PAGEREF {entry["anchor"]} \\h')


def add_clickable_toc(doc: Document, entries: list[dict[str, str | int]]) -> None:
    add_reference_list(doc, entries, "section")


def add_numbered_reference_list(doc: Document, entries: dict[int, dict[str, str | int]], label: str) -> None:
    for entry in entries.values():
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.45), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        title = f"{label} {entry['number']}. {entry['title']}"
        add_internal_hyperlink(paragraph, title, str(entry["anchor"]))
        paragraph.add_run("\t")
        add_field(paragraph, f'PAGEREF {entry["anchor"]} \\h')


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def apply_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(ACCENT)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.16

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 18, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, ACCENT, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(ACCENT)
        style.paragraph_format.space_after = Pt(4)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_para = section.header.paragraphs[0]
    header_para.text = "Software Engineering Specification - Multimodal Video Understanding Engine"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.size = Pt(9)
        run.font.name = "Aptos"
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer_para = section.footer.paragraphs[0]
    footer_para.text = ""
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run("Version 0.3 Draft | Page ")
    add_field(footer_para, "PAGE", "1")
    for run in footer_para.runs:
        run.font.size = Pt(9)
        run.font.name = "Aptos"
        run.font.color.rgb = RGBColor.from_string(MUTED)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip()
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        is_separator = all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)
        if not is_separator:
            rows.append(cells)
        index += 1
    return rows, index


def clear_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        paragraph._p.remove(run._r)


def add_markup_runs(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            link_match = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                add_external_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    for run in paragraph.runs:
        if not run.font.name:
            run.font.name = "Aptos"
        if not run.font.size:
            run.font.size = Pt(11)
        if run.font.color.rgb is None:
            run.font.color.rgb = RGBColor.from_string(ACCENT)


def add_caption(doc: Document, text: str, anchor: str, bookmark_id: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    add_bookmark(paragraph, anchor, bookmark_id)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(ACCENT)


def add_table(doc: Document, rows: list[list[str]], table_info: dict[str, str | int] | None = None) -> None:
    if not rows:
        return
    if table_info:
        add_caption(
            doc,
            f"Table {table_info['number']}. {table_info['title']}",
            str(table_info["anchor"]),
            int(table_info["bookmark_id"]),
        )
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)

    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        for col_index in range(col_count):
            cell = row.cells[col_index]
            text = row_data[col_index] if col_index < len(row_data) else ""
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_markup_runs(paragraph, text)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Aptos" if run.font.name != "Consolas" else "Consolas"
                    if run.font.color.rgb is None:
                        run.font.color.rgb = RGBColor.from_string(ACCENT)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    doc.add_paragraph()


def add_paragraph_with_inline_code(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    add_markup_runs(paragraph, text)
    return paragraph


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_example_block(doc: Document, language: str, code_lines: list[str]) -> None:
    title_by_language = {
        "json": "Structured response example",
        "http": "Request format example",
        "text": "Reference layout",
    }
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(4)
    label.paragraph_format.space_after = Pt(2)
    run = label.add_run(title_by_language.get(language, "Implementation example"))
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Aptos"
    run.font.color.rgb = RGBColor.from_string(MUTED)

    for line in code_lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.right_indent = Inches(0.05)
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.05
        set_paragraph_shading(paragraph, SOFT_FILL)
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(ACCENT)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_diagram_image(doc: Document, heading: str, figure_info: dict[str, str | int] | None = None) -> bool:
    filename = DIAGRAM_IMAGES.get(heading)
    if not filename:
        return False
    path = ROOT / "docs" / "diagrams" / filename
    if not path.exists():
        return False

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.4))

    if figure_info:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(10)
        add_bookmark(caption, str(figure_info["anchor"]), int(figure_info["bookmark_id"]))
        caption_run = caption.add_run(f"Figure {figure_info['number']}. {figure_info['title']}")
        caption_run.bold = True
        caption_run.font.size = Pt(9)
        caption_run.font.name = "Aptos"
        caption_run.font.color.rgb = RGBColor.from_string(MUTED)
    return True


def add_markdown_image(doc: Document, stripped: str, figure_info: dict[str, str | int] | None = None) -> bool:
    match = re.fullmatch(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)", stripped)
    if not match:
        return False
    rel_path = match.group("path")
    image_path = (ROOT / "docs" / rel_path).resolve()
    if not image_path.exists():
        return False

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.4))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    if figure_info:
        add_bookmark(caption, str(figure_info["anchor"]), int(figure_info["bookmark_id"]))
        caption_text = f"Figure {figure_info['number']}. {figure_info['title']}"
    else:
        caption_text = match.group("alt")
    caption_run = caption.add_run(caption_text)
    caption_run.bold = True
    caption_run.font.size = Pt(9)
    caption_run.font.name = "Aptos"
    caption_run.font.color.rgb = RGBColor.from_string(MUTED)
    return True


def build_docx() -> None:
    doc = Document()
    apply_document_styles(doc)
    add_header_footer(doc)
    set_update_fields_on_open(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    toc_entries = collect_toc_entries(lines)
    table_entries = collect_table_entries(lines)
    figure_entries = collect_figure_entries(lines)
    heading_queue = build_heading_queue(toc_entries)
    index = 0
    current_heading = ""
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if add_markdown_image(doc, stripped, figure_entries.get(index)):
            index += 1
            continue

        if stripped.startswith("|"):
            table_start = index
            rows, index = parse_table(lines, index)
            add_table(doc, rows, table_entries.get(table_start))
            continue

        if stripped == "## Table of Contents":
            paragraph = doc.add_heading("Table of Contents", level=1)
            add_clickable_toc(doc, toc_entries)
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("## "):
                index += 1
            continue

        if stripped == "## List of Tables":
            paragraph = doc.add_heading("List of Tables", level=1)
            if heading_queue[(1, "List of Tables")]:
                entry = heading_queue[(1, "List of Tables")].popleft()
                add_bookmark(paragraph, str(entry["anchor"]), int(entry["bookmark_id"]))
            add_numbered_reference_list(doc, table_entries, "Table")
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("## "):
                index += 1
            continue

        if stripped == "## List of Figures":
            paragraph = doc.add_heading("List of Figures", level=1)
            if heading_queue[(1, "List of Figures")]:
                entry = heading_queue[(1, "List of Figures")].popleft()
                add_bookmark(paragraph, str(entry["anchor"]), int(entry["bookmark_id"]))
            add_numbered_reference_list(doc, figure_entries, "Figure")
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("## "):
                index += 1
            continue

        if stripped.startswith("```"):
            language = stripped.removeprefix("```").strip()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip())
                index += 1
            if index < len(lines):
                index += 1
            if language == "mermaid" and add_diagram_image(doc, current_heading):
                continue
            add_example_block(doc, language, code_lines)
            continue

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            run.bold = True
            run.font.size = Pt(22)
            run.font.name = "Aptos Display"
            run.font.color.rgb = RGBColor.from_string(ACCENT)
        elif stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            if heading_text.startswith("Part "):
                doc.add_page_break()
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(14)
                run = paragraph.add_run(heading_text)
                run.bold = True
                run.font.name = "Aptos Display"
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor.from_string(ACCENT)
            else:
                paragraph = doc.add_heading(heading_text, level=1)
            if heading_queue[(1, heading_text)]:
                entry = heading_queue[(1, heading_text)].popleft()
                add_bookmark(paragraph, str(entry["anchor"]), int(entry["bookmark_id"]))
        elif stripped.startswith("### "):
            current_heading = stripped[4:].strip()
            paragraph = doc.add_heading(current_heading, level=2)
            if heading_queue[(2, current_heading)]:
                entry = heading_queue[(2, current_heading)].popleft()
                add_bookmark(paragraph, str(entry["anchor"]), int(entry["bookmark_id"]))
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
        elif stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(stripped[2:].strip())
            paragraph.paragraph_format.space_after = Pt(4)
        elif re.match(r"^\d+\.\s+", stripped):
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.add_run(re.sub(r"^\d+\.\s+", "", stripped))
            paragraph.paragraph_format.space_after = Pt(4)
        else:
            add_paragraph_with_inline_code(doc, stripped)

        index += 1

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
